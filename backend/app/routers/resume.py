from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from typing import Dict, Optional
import PyPDF2
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
import json
import difflib
import base64

from ..database import get_db
from ..models import User
from .users import get_current_user
from app.config import settings
from openai import OpenAI

router = APIRouter(prefix="/api/resume", tags=["resume"])

_openai = OpenAI(api_key=settings.openai_api_key)


def extract_text_from_resume(file_content: bytes, filename: str) -> str:
    """
    Extract text from resume file (PDF or DOCX)
    Returns: extracted text as string
    """
    try:
        if filename.endswith('.pdf'):
            # Parse PDF
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        
        elif filename.endswith('.docx'):
            # Parse DOCX
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            return text
        
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format")
    
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to read file: {str(e)}")


def parse_resume_text(text: str) -> Dict[str, any]:
    """
    Parse resume text and extract relevant information
    Returns: dict with extracted fields
    """
    try:
        # Initialize result
        result = {
            "full_name": None,
            "target_role": None,
            "skills": [],
            "location_preference": None,
        }
        
        # Extract name (usually in first few lines, capitalized)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if lines:
            # First line is often the name
            potential_name = lines[0]
            if len(potential_name.split()) <= 4 and potential_name[0].isupper():
                result["full_name"] = potential_name
        
        # Extract email (for validation)
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        
        # Extract job titles (common patterns)
        title_keywords = [
            "Software Engineer", "Developer", "Data Scientist", "Product Manager",
            "Designer", "Analyst", "Marketing", "Sales", "Manager", "Director",
            "Frontend", "Backend", "Full Stack", "DevOps", "ML Engineer",
            "UX Designer", "UI Designer", "Project Manager", "QA Engineer"
        ]
        
        text_lower = text.lower()
        for keyword in title_keywords:
            if keyword.lower() in text_lower:
                result["target_role"] = keyword
                break
        
        # Extract skills (common tech skills)
        skill_patterns = [
            # Programming languages
            r'\b(Python|JavaScript|Java|C\+\+|C#|Ruby|PHP|Swift|Kotlin|Go|Rust|TypeScript)\b',
            # Frameworks/Libraries
            r'\b(React|Angular|Vue|Node\.?js|Django|Flask|Spring|\.NET|Laravel|Rails)\b',
            # Databases
            r'\b(SQL|MySQL|PostgreSQL|MongoDB|Redis|Oracle|SQLite)\b',
            # Cloud/DevOps
            r'\b(AWS|Azure|GCP|Docker|Kubernetes|Jenkins|Git|CI/CD)\b',
            # Other tech
            r'\b(HTML|CSS|REST|GraphQL|API|Agile|Scrum|Machine Learning|AI|Data Analysis)\b',
        ]
        
        skills_set = set()
        for pattern in skill_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                # Normalize skill name
                skill = match.strip()
                if skill:
                    skills_set.add(skill)
        
        result["skills"] = list(skills_set)[:20]  # Limit to 20 skills
        
        # Extract location (cities/states pattern)
        location_pattern = r'\b(New York|San Francisco|Los Angeles|Chicago|Boston|Seattle|Austin|Denver|Remote|USA|United States)\b'
        locations = re.findall(location_pattern, text, re.IGNORECASE)
        if locations:
            result["location_preference"] = locations[0]
        
        return result
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse resume: {str(e)}")


@router.post("/upload")
async def upload_resume(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Upload and parse resume (PDF or DOCX), return extracted information
    """
    # Validate file type
    if not (file.filename.endswith('.pdf') or file.filename.endswith('.docx')):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
    
    # Validate file size (max 5MB)
    content = await file.read()
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File size must be less than 5MB")
    
    # Extract text from file
    text = extract_text_from_resume(content, file.filename)
    
    # Parse resume text
    parsed_data = parse_resume_text(text)
    
    # Update user profile with parsed data (if fields are empty)
    user = db.query(User).filter(User.id == current_user.id).first()
    
    updated_fields = []
    
    if parsed_data["full_name"] and not user.full_name:
        user.full_name = parsed_data["full_name"]
        updated_fields.append("full_name")
    
    if parsed_data["target_role"] and not user.target_role:
        user.target_role = parsed_data["target_role"]
        updated_fields.append("target_role")
    
    if parsed_data["skills"] and not user.skills:
        user.skills_list = parsed_data["skills"]
        updated_fields.append("skills")
    
    if parsed_data["location_preference"] and not user.location_preference:
        user.location_preference = parsed_data["location_preference"]
        updated_fields.append("location_preference")
    
    if updated_fields:
        db.commit()
        db.refresh(user)

    return {
        "message": "Resume parsed successfully",
        "parsed_data": parsed_data,
        "updated_fields": updated_fields,
        "user": {
            "full_name": user.full_name,
            "target_role": user.target_role,
            "skills": user.skills_list,
            "location_preference": user.location_preference,
        }
    }


# ---------------------------------------------------------------------------
# Helpers for resume rewrite
# ---------------------------------------------------------------------------

def _extract_docx_structure(content: bytes):
    """
    Extract paragraph structure from a DOCX file.
    Returns (plain_text, list of {text, style_name} dicts).
    """
    doc = docx.Document(io.BytesIO(content))
    paragraphs = []
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append({"text": para.text.strip(), "style": para.style.name})
            lines.append(para.text.strip())
    # Also grab table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append({"text": cell.text.strip(), "style": "Table Contents"})
                    lines.append(cell.text.strip())
    return "\n".join(lines), paragraphs


async def _rewrite_with_ai(resume_text: str, job_description: str, extra_context: str = "") -> list:
    """
    Call GPT-4o-mini to rewrite the resume for the job description.
    Returns a list of section dicts: [{heading, level, lines}].
    extra_context: confirmed Q&A answers from the candidate — must be added to the resume.
    """
    confirmed_additions = ""
    confirmed_rule = "Only rephrase, reorder, and re-emphasize content already in the resume."
    if extra_context and extra_context.strip():
        confirmed_additions = f"""
=== CONFIRMED ADDITIONS (candidate verified — MUST be added to the resume) ===
{extra_context.strip()}
=== END CONFIRMED ADDITIONS ===

"""
        confirmed_rule = (
            "You MUST incorporate every item from CONFIRMED ADDITIONS into the resume. "
            "Add them to the Skills section AND add bullet points in the most relevant Experience entry."
        )

    prompt = f"""You are a professional resume writer rewriting a resume to match a job description.
{confirmed_additions}
RULES:
1. NEVER fabricate job titles, company names, dates, or achievements not in the resume or CONFIRMED ADDITIONS.
2. {confirmed_rule}
3. Use keywords from the job description only where they match the candidate's real experience.
4. Keep ALL original sections (Education, Experience, Skills, etc.) in the same order.
5. Do not remove any real experience or skills from the resume.

ORIGINAL RESUME:
{resume_text}

JOB DESCRIPTION:
{job_description[:3000]}

Return a JSON object with this exact structure:
{{
  "sections": [
    {{
      "heading": "section heading or empty string for sub-content",
      "level": 0,
      "lines": ["line 1", "line 2", ...]
    }}
  ]
}}

Guidelines for the JSON:
- level 1 = candidate name (first section)
- level 2 = major section headings (Experience, Education, Skills, Summary, etc.)
- level 0 = body content lines (job entries, descriptions, skill lists)
- Use "• " prefix for bullet points
- Keep each line concise (one idea per line)"""

    try:
        response = _openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are an expert resume writer who tailors resumes to job descriptions. "
                        "You work with content from the original resume and any supplemental "
                        "information provided by the candidate. You never fabricate facts."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
            max_tokens=3000,
            response_format={"type": "json_object"},
        )
        data = json.loads(response.choices[0].message.content)
        return data.get("sections", [])
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI rewrite failed: {str(e)}")


def _sections_to_text(sections: list) -> str:
    """Flatten sections back to plain text for the evaluation pass."""
    lines = []
    for sec in sections:
        heading = sec.get("heading", "").strip()
        if heading:
            lines.append(f"\n{heading}")
        for line in sec.get("lines", []):
            if isinstance(line, dict):
                sub_heading = line.get("heading", "")
                if sub_heading:
                    lines.append(sub_heading)
                for sub in line.get("lines", []):
                    if isinstance(sub, str):
                        lines.append(f"  {sub}")
            elif isinstance(line, str):
                lines.append(line)
    return "\n".join(lines)


async def _evaluate_and_correct(sections: list, job_description: str, extra_context: str = "") -> list:
    """
    Second LLM pass: HR specialist reviews the rewritten resume against the
    job description and corrects any issues (weak verbs, vague bullets, poor
    alignment, formatting inconsistencies, etc.) without inventing new content.
    Returns corrected sections in the same JSON structure.
    """
    resume_text = _sections_to_text(sections)

    confirmed_block = ""
    if extra_context and extra_context.strip():
        confirmed_block = f"""
CONFIRMED ADDITIONS — if any of the items below are missing from the draft resume, you MUST add them:
{extra_context.strip()}

"""

    prompt = f"""You are a senior HR specialist and certified professional resume coach with 15+ years of experience screening and evaluating resumes for tech roles.

You have been given a draft resume that was tailored to a specific job description. Your task is to EVALUATE it and CORRECT any issues.
{confirmed_block}
CHECK FOR AND FIX:
1. Weak or passive verbs in bullet points → replace with strong action verbs (Led, Built, Optimized, Delivered, Reduced, Increased, Designed, etc.)
2. Vague or generic bullet points → make them specific and achievement-oriented where evidence exists in the text
3. Poor alignment with the job description → reorder or re-emphasize content to front-load the most relevant experience
4. Missing quantification opportunities → add numbers/metrics ONLY if they are already mentioned in the resume
5. Redundant or filler phrases → remove them
6. Inconsistent formatting (tense, punctuation, style) → standardize throughout
7. Summary/objective not tailored to the role → improve it based on the candidate's actual background
8. Skills section not optimally ordered → put the most JD-relevant skills first

ABSOLUTE RULES — never break these:
- NEVER add skills, technologies, certifications, or experiences NOT in the draft resume (CONFIRMED ADDITIONS above are allowed)
- NEVER fabricate metrics, company names, dates, or achievements
- NEVER remove legitimate experience or skills
- Keep the same sections in the same order

DRAFT RESUME:
{resume_text}

JOB DESCRIPTION:
{job_description[:3000]}

Return the corrected resume as a JSON object with the EXACT same structure as the input:
{{
  "sections": [
    {{
      "heading": "section heading or empty string",
      "level": 1 or 2 or 0,
      "lines": ["line 1", "line 2", ...]
    }}
  ]
}}

IMPORTANT: lines must be plain strings only — no nested objects."""

    try:
        response = _openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a senior HR specialist and certified professional resume coach. "
                        "You evaluate and improve resumes without fabricating any information. "
                        "You only improve language, impact, and alignment — never the facts."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=2500,
            response_format={"type": "json_object"},
        )
        data = json.loads(response.choices[0].message.content)
        corrected = data.get("sections", [])
        # Fall back to original if evaluation returned nothing
        return corrected if corrected else sections
    except Exception:
        # Evaluation is best-effort — return original sections on failure
        return sections


def _build_rewritten_docx(sections: list) -> bytes:
    """
    Build a well-formatted DOCX from the rewritten sections.
    """
    out = docx.Document()

    # Set reasonable margins
    for section in out.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    for sec in sections:
        heading = sec.get("heading", "").strip()
        level = sec.get("level", 0)
        lines = sec.get("lines", [])

        if heading:
            if level == 1:
                # Candidate name — large, bold, centered
                p = out.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(heading)
                run.bold = True
                run.font.size = Pt(18)
            elif level == 2:
                # Section heading
                h = out.add_heading(heading, level=2)
                h.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)  # indigo-900
            else:
                p = out.add_paragraph()
                p.add_run(heading).bold = True

        for line in lines:
            # AI sometimes nests sub-sections (e.g. job entries) as dicts inside lines
            if isinstance(line, dict):
                sub_heading = line.get("heading", "").strip()
                sub_lines = line.get("lines", [])
                if sub_heading:
                    p = out.add_paragraph()
                    p.add_run(sub_heading).bold = True
                for sub_line in sub_lines:
                    if isinstance(sub_line, dict):
                        sub_line = sub_line.get("text") or sub_line.get("line") or ""
                    sub_line = sub_line.strip()
                    if not sub_line:
                        continue
                    if sub_line.startswith("• ") or sub_line.startswith("- "):
                        out.add_paragraph(sub_line.lstrip("•- ").strip(), style="List Bullet")
                    else:
                        out.add_paragraph(sub_line)
                continue
            line = line.strip()
            if not line:
                continue
            if line.startswith("• ") or line.startswith("- "):
                out.add_paragraph(line.lstrip("•- ").strip(), style="List Bullet")
            else:
                out.add_paragraph(line)

    buf = io.BytesIO()
    out.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Rewrite endpoint
# ---------------------------------------------------------------------------

@router.post("/rewrite")
async def rewrite_resume(
    file: UploadFile = File(...),
    job_description: str = Form(...),
    current_user: User = Depends(get_current_user),
):
    """
    Rewrite an uploaded resume (PDF or DOCX) tailored to the given job description.
    Returns a .docx file preserving the original resume structure.
    Only uses skills and experience present in the original resume — nothing is fabricated.
    """
    if not (file.filename.endswith(".pdf") or file.filename.endswith(".docx")):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")

    content = await file.read()
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File size must be less than 5MB")

    if not job_description.strip():
        raise HTTPException(status_code=400, detail="Job description is required")

    # Extract resume text
    if file.filename.endswith(".docx"):
        resume_text, _ = _extract_docx_structure(content)
    else:
        resume_text = extract_text_from_resume(content, file.filename)

    # Pass 1: rewrite the resume for the job description
    sections = await _rewrite_with_ai(resume_text, job_description, "")

    # Pass 2: HR specialist evaluation and correction (behind the scenes)
    sections = await _evaluate_and_correct(sections, job_description, "")

    # Build output DOCX
    docx_bytes = _build_rewritten_docx(sections)

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": 'attachment; filename="rewritten_resume.docx"'},
    )


@router.post("/analyze-gaps")
async def analyze_resume_gaps(
    file: UploadFile = File(...),
    job_description: str = Form(...),
    current_user: User = Depends(get_current_user),
):
    """
    Compare the candidate's resume against the job description.
    Returns missing requirements and targeted questions to uncover hidden experience.
    Response: { summary: str, gaps: [{requirement: str, question: str}] }
    """
    if not (file.filename.endswith(".pdf") or file.filename.endswith(".docx")):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")

    content = await file.read()
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File size must be less than 5MB")

    if not job_description.strip():
        raise HTTPException(status_code=400, detail="Job description is required")

    if file.filename.endswith(".docx"):
        resume_text, _ = _extract_docx_structure(content)
    else:
        resume_text = extract_text_from_resume(content, file.filename)

    prompt = f"""You are a resume gaps analyst. Compare the candidate's resume against the job description to identify what requirements are missing or weakly represented.

For each gap, write a short, specific question to ask the candidate — they may have the experience but forgot to include it on their resume.

RESUME:
{resume_text[:3000]}

JOB DESCRIPTION:
{job_description[:3000]}

Return JSON with this exact structure:
{{
  "summary": "1-2 sentence overview of the key gaps between this resume and the job requirements",
  "gaps": [
    {{
      "requirement": "Short label for the missing requirement (e.g. 'Kubernetes experience')",
      "question": "Specific question to ask the candidate (e.g. 'Have you worked with Kubernetes or container orchestration in any of your roles?')"
    }}
  ]
}}

Return at most 5 gaps. Focus on the most important gaps that would affect getting an interview."""

    try:
        response = _openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a professional resume gaps analyst. "
                        "You identify missing requirements and generate targeted questions "
                        "to help candidates surface relevant experience they may have forgotten to include."
                    )
                },
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=800,
            response_format={"type": "json_object"},
        )
        data = json.loads(response.choices[0].message.content)
        return {
            "summary": data.get("summary", ""),
            "gaps": data.get("gaps", []),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gap analysis failed: {str(e)}")


@router.post("/rewrite-diff")
async def rewrite_resume_diff(
    file: UploadFile = File(...),
    job_description: str = Form(...),
    extra_context: Optional[str] = Form(None),
    current_user: User = Depends(get_current_user),
):
    """
    Same 2-pass rewrite as /rewrite but returns JSON instead of a file:
    - diff: line-by-line git-style diff (added / removed / context)
    - docx_b64: base64-encoded .docx of the final approved version
    extra_context: optional candidate Q&A answers from gap analysis to enrich the rewrite.
    """
    if not (file.filename.endswith(".pdf") or file.filename.endswith(".docx")):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")

    content = await file.read()
    if len(content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File size must be less than 5MB")

    if not job_description.strip():
        raise HTTPException(status_code=400, detail="Job description is required")

    # Extract original text
    if file.filename.endswith(".docx"):
        original_text, _ = _extract_docx_structure(content)
    else:
        original_text = extract_text_from_resume(content, file.filename)

    # 2-pass rewrite (extra_context from gap Q&A included if provided)
    sections = await _rewrite_with_ai(original_text, job_description, extra_context or "")
    sections = await _evaluate_and_correct(sections, job_description, extra_context or "")
    rewritten_text = _sections_to_text(sections)

    # Compute diff
    orig_lines = [l for l in original_text.splitlines() if l.strip()]
    new_lines  = [l for l in rewritten_text.splitlines() if l.strip()]

    diff_lines = []
    for raw in difflib.ndiff(orig_lines, new_lines):
        code = raw[:2]
        text = raw[2:].strip()
        if not text:
            continue
        if code == "+ ":
            diff_lines.append({"type": "added",   "text": text})
        elif code == "- ":
            diff_lines.append({"type": "removed",  "text": text})
        elif code == "  ":
            diff_lines.append({"type": "context",  "text": text})
        # skip '? ' annotation lines

    # Build DOCX and base64-encode it for download after approval
    docx_bytes = _build_rewritten_docx(sections)
    docx_b64 = base64.b64encode(docx_bytes).decode()

    return {"diff": diff_lines, "docx_b64": docx_b64}
